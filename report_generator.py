from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from io import BytesIO
import pandas as pd
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import unicodedata

# ---------------- Nombre de archivo ----------------
def _slug_empresa(nombre: str) -> str:
    nfkd = unicodedata.normalize("NFKD", nombre)
    ascii_only = nfkd.encode("ascii", "ignore").decode("ascii")
    return "".join(ch for ch in ascii_only if ch.isalnum())

def build_report_filename(empresa: str, date: datetime | None = None) -> str:
    date = date or datetime.now()
    date_tag = date.strftime("%y%m%d")  # YYMMDD
    empresa_tag = _slug_empresa(empresa)
    return f"{date_tag}-LV-{empresa_tag}-Facturación honorarios.docx"

# --- Colores y Fuentes del Diseño ---
FONT_FAMILY = 'Calibri Light'
COLOR_PRIMARY = RGBColor(0, 51, 102)   # #003366
COLOR_ACCENT  = RGBColor(226, 0, 116)  # #E20074
COLOR_LIGHT_GRAY = RGBColor(240, 240, 240)
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_WHITE = RGBColor(255, 255, 255)

# Agregar estas funciones directamente en el archivo:
def format_currency(value, currency="USD"):
    """Formatea un número como moneda."""
    try:
        return f"{currency} {float(value):,.2f}"
    except (ValueError, TypeError):
        return f"{currency} 0.00"

def find_column(df, possible_names):
    """Busca una columna en el DataFrame usando una lista de posibles nombres."""
    for col_name in possible_names:
        for actual_col in df.columns:
            if col_name.upper() in actual_col.upper():
                return actual_col
    return None

def get_document_count(df):
    """Obtiene el número de documentos únicos."""
    possible_names = ['NO. CASO', 'NUMERO CASO', 'CASO', 'ID', 'NUMERO', 'DOCUMENTO']
    col_name = find_column(df, possible_names)
    
    if col_name:
        return df[col_name].nunique()
    else:
        return len(df)

# -------------------------------------------------
# Utilidades de XML / Bordes
# -------------------------------------------------
def set_table_borders(table, color="000000", size=8, val="single"):
    """Bordes de tabla con interiores."""
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.append(tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV'):
        el = OxmlElement(edge); el.set(qn('w:val'), val); el.set(qn('w:sz'), str(int(size)))
        el.set(qn('w:color'), color); el.set(qn('w:space'), '0'); tblBorders.append(el)
    tblPr.append(tblBorders)

def set_table_outer_borders(table, color="000000", size=8, val="single"):
    """Solo bordes exteriores (sin insideH/insideV)."""
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.append(tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('w:top', 'w:left', 'w:bottom', 'w:right'):
        el = OxmlElement(edge); el.set(qn('w:val'), val); el.set(qn('w:sz'), str(int(size)))
        el.set(qn('w:color'), color); el.set(qn('w:space'), '0'); tblBorders.append(el)
    tblPr.append(tblBorders)

def set_cell_borders(cell, left=False, right=False, top=False, bottom=False, color="000000", size=8, val="single"):
    """Bordes individuales por celda (override de los de tabla)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    def _edge(tag, on):
        # elimina si existe
        old = borders.find(qn(f'w:{tag}'))
        if old is not None: borders.remove(old)
        if on:
            el = OxmlElement(f'w:{tag}')
            el.set(qn('w:val'), val); el.set(qn('w:sz'), str(int(size)))
            el.set(qn('w:color'), color); el.set(qn('w:space'), '0')
            borders.append(el)
    _edge('left', left); _edge('right', right); _edge('top', top); _edge('bottom', bottom)

def _to_hex(color) -> str:
    if isinstance(color, RGBColor): return str(color)
    if isinstance(color, (tuple, list)) and len(color) == 3:
        r, g, b = map(int, color); return f"{r:02x}{g:02x}{b:02x}"
    if isinstance(color, str):
        s = color.lstrip("#"); return s.upper() if len(s) == 6 else "FFFFFF"
    return "FFFFFF"

def set_cell_shading(cell, color_rgb):
    hex_color = _to_hex(color_rgb)
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def set_paragraph_border_bottom(paragraph, color="000000", size=6, space=1):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pb = pPr.find(qn('w:pBdr'))
    if pb is not None: pPr.remove(pb)
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), str(space)); bottom.set(qn('w:color'), color)
    pb.append(bottom); pPr.append(pb)

# -------------------------------------------------
# Estilos y utilidades varias
# -------------------------------------------------
def style_table(table, has_total_row=True):
    for i, row in enumerate(table.rows):
        is_header = (i == 0)
        is_total  = has_total_row and (i == len(table.rows) - 1)
        for cell in row.cells:
            if is_header:
                set_cell_shading(cell, COLOR_PRIMARY)
            elif is_total:
                set_cell_shading(cell, COLOR_ACCENT)
            else:
                set_cell_shading(cell, COLOR_LIGHT_GRAY)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1
                for run in p.runs:
                    run.font.name = FONT_FAMILY; run.font.size = Pt(11)
                    if is_header or is_total:
                        run.font.color.rgb = COLOR_WHITE; run.font.bold = True
                    else:
                        run.font.color.rgb = COLOR_BLACK; run.font.bold = False
def set_row_text_color(row, color_rgb=COLOR_WHITE, bold=True):
    """Pone el color de fuente (y negrita opcional) a todos los runs de una fila."""
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = color_rgb
                if bold is not None:
                    run.bold = bold

def get_representative_price(data: pd.DataFrame) -> float:
    if 'VALOR' not in data.columns: return 0.0
    serie = pd.to_numeric(data['VALOR'], errors='coerce').dropna()
    if serie.empty: return 0.0
    moda = serie.mode()
    return float(moda.iloc[0]) if not moda.empty else float(serie.iloc[0])

def merge_row_cells(table, row_idx: int, start_col: int, end_col: int):
    top_left = table.cell(row_idx, start_col)
    bottom_right = table.cell(row_idx, end_col)
    return top_left.merge(bottom_right)

def _fix_table_layout_3cols(table, widths=(Inches(1.2), Inches(4.2), Inches(1.4))):
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1

def _set_row_height(row, height_in_inches=0.28):
    row.height = Inches(height_in_inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

def apply_3col_grid_borders(table, skip_mid_rows=()):
    """
    Dibuja bordes por celda en una tabla de 3 columnas:
    - Bordes exteriores en todo el perímetro
    - Separador entre col0-col1 salvo en filas de skip_mid_rows (para filas fusionadas)
    - Separador entre col1-col2 en todas las filas
    """
    rows = list(range(len(table.rows)))
    last_row = rows[-1]
    for r in rows:
        # top/bottom para toda la fila
        for c in range(3):
            left = (c == 0)
            right = (c == 2)
            top = (r == 0)
            bottom = (r == last_row)
            set_cell_borders(table.cell(r, c), left=left, right=right, top=top, bottom=bottom)
        # separador entre col0 y col1 (evítalo si la fila está fusionada)
        if r not in skip_mid_rows:
            set_cell_borders(table.cell(r, 0), right=True)
            set_cell_borders(table.cell(r, 1), left=True)
        # separador entre col1 y col2 (siempre)
        set_cell_borders(table.cell(r, 1), right=True)
        set_cell_borders(table.cell(r, 2), left=True)

# -------------------------------------------------
# Tablas del documento
# -------------------------------------------------
def add_main_table(doc, data, empresa: str):
    """Añade la tabla principal de datos al documento (con merge en fila Total)."""
    doc.add_paragraph()

    required_cols = ['MES ASIGNACION', 'AÑO ASIGNACION', 'NOMBRE', 'MONEDA', 'VALOR']
    available_cols = [col for col in required_cols if col in data.columns]
    if not available_cols:
        doc.add_paragraph("Error: No se encontraron las columnas necesarias en los datos.")
        return

    df_main = data[available_cols].copy()
    table = doc.add_table(rows=1, cols=len(available_cols))
    table.autofit = True

    # Encabezados
    for i, name in enumerate(available_cols):
        table.cell(0, i).text = name

    # Filas
    for _, row in df_main.iterrows():
        cells = table.add_row().cells
        for i, col_name in enumerate(available_cols):
            value = row[col_name]
            if col_name == 'VALOR':
                try: cells[i].text = f"{float(value):,.2f}"
                except Exception: cells[i].text = str(value)
            else:
                cells[i].text = str(value)

    # Fila Total con merge horizontal (todo excepto 'VALOR')
    if 'VALOR' in available_cols:
        val_idx = available_cols.index('VALOR')
        table.add_row()
        total_row_idx = len(table.rows) - 1

        merged_cell = merge_row_cells(table, total_row_idx, 0, max(0, val_idx - 1))
        label = "Total (precio único)" if empresa == "Gwealth" else "Total"
        merged_cell.text = label  # <-- escribir DESPUÉS de fusionar (evita salto)
        for p in merged_cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        # Valor en la columna 'VALOR'
        if empresa == "Gwealth":
            total_unique = get_representative_price(data)
            table.cell(total_row_idx, val_idx).text = f"{total_unique:,.2f}"
        else:
            table.cell(total_row_idx, val_idx).text = f"{data['VALOR'].sum():,.2f}"

    style_table(table)
    set_table_borders(table)  # aquí sí queremos interiores

def add_summary_tables(doc, data, empresa, anio, mes):
    """Añade las tablas de resumen específicas por empresa (excluyendo Ravago)."""
    total_valor_sum = data['VALOR'].sum() if 'VALOR' in data.columns else 0.0
    doc.add_paragraph()

    if empresa == "Altimetrik":
        table = doc.add_table(rows=2, cols=3)
        table.cell(0,0).text = "Mes"
        table.cell(0,1).text = "Concepto"
        table.cell(0,2).text = "Total"
        table.cell(1,0).text = mes
        table.cell(1,1).text = f"Consultas en listas recibidas en {mes} de {anio}"
        table.cell(1,2).text = format_currency(total_valor_sum)

        style_table(table, has_total_row=False)
        _fix_table_layout_3cols(table)
        for r in table.rows: _set_row_height(r, 0.28)
        set_table_borders(table)

    elif empresa == "Gwealth":
        precio_unico = get_representative_price(data)
        iva = precio_unico * 0.19
        total_con_iva = precio_unico + iva

        table = doc.add_table(rows=4, cols=3)
        # Encabezados
        table.cell(0,0).text = "Mes"
        table.cell(0,1).text = "Concepto"
        table.cell(0,2).text = "Total"

        # Fila de contenido
        table.cell(1,0).text = mes
        table.cell(1,1).text = f"Consultas en listas recibidas en {mes} de {anio}"
        table.cell(1,2).text = format_currency(precio_unico)

        # Fila TOTAL (combinar col 0 y 1) -> escribir texto DESPUÉS de fusionar
        table.cell(2,0).text = ""
        table.cell(2,1).text = ""
        merged_total = merge_row_cells(table, 2, 0, 1)
        merged_total.text = "TOTAL"
        for p in merged_total.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        table.cell(2,2).text = format_currency(precio_unico)

        # Fila TOTAL CON IVA (igual)
        table.cell(3,0).text = ""
        table.cell(3,1).text = ""
        merged_total_iva = merge_row_cells(table, 3, 0, 1)
        merged_total_iva.text = "TOTAL CON IVA"
        for p in merged_total_iva.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        table.cell(3,2).text = format_currency(total_con_iva)

        # Estilos + colores (las filas 2 y 3 en acento)
        style_table(table, has_total_row=False)
        for cell in table.rows[2].cells: set_cell_shading(cell, COLOR_ACCENT)
        for cell in table.rows[3].cells: set_cell_shading(cell, COLOR_ACCENT)
        set_row_text_color(table.rows[2], COLOR_WHITE, bold=True)
        set_row_text_color(table.rows[3], COLOR_WHITE, bold=True)

        # Layout compacto y consistente
        _fix_table_layout_3cols(table, widths=(Inches(1.2), Inches(4.2), Inches(1.4)))
        for r in table.rows: _set_row_height(r, 0.28)

        # Bordes: solo exteriores a nivel tabla + rejilla por celda (sin línea en filas fusionadas)
        set_table_outer_borders(table)  # quita insideH/insideV
        apply_3col_grid_borders(table, skip_mid_rows=(2, 3))  # no dibujar separador 0-1 en filas 2 y 3

# -------------------------------------------------
# Generación del documento
# -------------------------------------------------
def generate_report(data, empresa, anio, mes, funcionarios):
    """Genera el documento Word desde cero para Altimetrik y GWealth."""
    doc = Document()

    normal_style = doc.styles['Normal']
    normal_style.font.name = FONT_FAMILY
    normal_style.font.size = Pt(11)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Logo en cuerpo
    p_logo = doc.add_paragraph(); p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        p_logo.add_run().add_picture('assets/biu_logo.png', width=Inches(1.5))
    except FileNotFoundError:
        p_logo.add_run("[Logo BIU]").bold = True

    # Regla
    p_rule = doc.add_paragraph()
    set_paragraph_border_bottom(p_rule, color="000000", size=6, space=1)

    # Títulos
    p = doc.add_paragraph(); run = p.add_run(f"FACTURACIÓN {mes.upper()} {anio}")
    run.font.name = FONT_FAMILY; run.font.size = Pt(24); run.bold = True; run.font.color.rgb = COLOR_PRIMARY
    p = doc.add_paragraph(); run = p.add_run(empresa.upper())
    run.font.name = FONT_FAMILY; run.font.size = Pt(20); run.bold = True; run.font.color.rgb = COLOR_PRIMARY

    # Info
    doc.add_paragraph(" \nFecha de corte del reporte: ")
    doc.add_paragraph(f"Funcionario que reporta: \t {funcionarios['reporta']}")
    doc.add_paragraph(f"Funcionario revisor: \t\t {funcionarios['revisor']}")

    # Tablas
    add_main_table(doc, data, empresa)
    add_summary_tables(doc, data, empresa, anio, mes)

    # Footer
    footer = doc.sections[0].footer
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.text = "Número: 601 - 7455289 | Dirección: Carrera 7 No. 74B-56, Oficina 301 | Correo: info@biu.com.co"
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_footer.runs:
        r.font.name = FONT_FAMILY; r.font.size = Pt(9)

    buffer = BytesIO()
    doc.save(buffer); buffer.seek(0)
    return buffer
