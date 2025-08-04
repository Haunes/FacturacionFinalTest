from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import pandas as pd
from utils import format_currency

def get_representative_price(data: pd.DataFrame) -> float:
    """Precio representativo para GWealth."""
    if 'VALOR' not in data.columns: 
        return 0.0
    serie = pd.to_numeric(data['VALOR'], errors='coerce').dropna()
    if serie.empty: 
        return 0.0
    moda = serie.mode()
    return float(moda.iloc[0]) if not moda.empty else float(serie.iloc[0])

from .word_table_styles import WordTableStyles

class WordTableBuilder:
    """Construye tablas para documentos Word."""
    
    def __init__(self):
        self.table_styles = WordTableStyles()
    
    def add_main_table(self, doc: Document, data: pd.DataFrame, empresa: str):
        """Añade la tabla principal de datos al documento."""
        required_cols = ['MES ASIGNACION', 'AÑO ASIGNACION', 'NOMBRE', 'MONEDA', 'VALOR']
        available_cols = [col for col in required_cols if col in data.columns]
        
        if not available_cols:
            doc.add_paragraph("Error: No se encontraron las columnas necesarias en los datos.")
            return

        df_main = data[available_cols].copy()
        table = doc.add_table(rows=1, cols=len(available_cols))
        table.autofit = True

        # Encabezados
        for i, name in enumerate(available_cols):
            table.cell(0, i).text = name

        # Filas de datos
        for _, row in df_main.iterrows():
            cells = table.add_row().cells
            for i, col_name in enumerate(available_cols):
                value = row[col_name]
                if col_name == 'VALOR':
                    try:
                        cells[i].text = f"{float(value):,.2f}"
                    except Exception:
                        cells[i].text = str(value)
                else:
                    cells[i].text = str(value)

        # Fila Total
        self._add_total_row(table, data, empresa, available_cols)
        
        # Aplicar estilos
        self.table_styles.style_table(table)
        self.table_styles.set_table_borders(table)

    def _add_total_row(self, table, data: pd.DataFrame, empresa: str, available_cols: list):
        """Agrega la fila de total a la tabla principal."""
        if 'VALOR' not in available_cols:
            return
        
        val_idx = available_cols.index('VALOR')
        table.add_row()
        total_row_idx = len(table.rows) - 1

        # Fusionar celdas para el label
        merged_cell = self.table_styles.merge_row_cells(table, total_row_idx, 0, max(0, val_idx - 1))
        label = "Total (precio único)" if empresa == "Gwealth" else "Total"
        merged_cell.text = label
        
        # Centrar el texto
        for p in merged_cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        # Valor total
        if empresa == "Gwealth":
            total_unique = get_representative_price(data)
            table.cell(total_row_idx, val_idx).text = f"{total_unique:,.2f}"
        else:
            table.cell(total_row_idx, val_idx).text = f"{data['VALOR'].sum():,.2f}"

    def add_summary_tables(self, doc: Document, data: pd.DataFrame, empresa: str, anio: int, mes: str):
        """Añade las tablas de resumen específicas por empresa."""
        if empresa == "Altimetrik":
            self._add_altimetrik_table(doc, data, anio, mes)
        elif empresa == "Gwealth":
            self._add_gwealth_table(doc, data, anio, mes)

    def _add_altimetrik_table(self, doc: Document, data: pd.DataFrame, anio: int, mes: str):
        """Agrega tabla específica para Altimetrik."""
        total_valor_sum = data['VALOR'].sum() if 'VALOR' in data.columns else 0.0
        
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "Mes"
        table.cell(0, 1).text = "Concepto"
        table.cell(0, 2).text = "Total"
        table.cell(1, 0).text = mes
        table.cell(1, 1).text = f"Consultas en listas recibidas en {mes} de {anio}"
        table.cell(1, 2).text = format_currency(total_valor_sum)

        self.table_styles.style_table(table, has_total_row=False)
        self.table_styles.fix_table_layout_3cols(table)
        self.table_styles.set_table_borders(table)

    def _add_gwealth_table(self, doc: Document, data: pd.DataFrame, anio: int, mes: str):
        """Agrega tabla específica para Gwealth."""
        precio_unico = get_representative_price(data)
        iva = precio_unico * 0.19
        total_con_iva = precio_unico + iva

        table = doc.add_table(rows=4, cols=3)
        
        # Encabezados
        table.cell(0, 0).text = "Mes"
        table.cell(0, 1).text = "Concepto"
        table.cell(0, 2).text = "Total"

        # Fila de contenido
        table.cell(1, 0).text = mes
        table.cell(1, 1).text = f"Consultas en listas recibidas en {mes} de {anio}"
        table.cell(1, 2).text = format_currency(precio_unico)

        # Filas de totales con fusión
        self._add_gwealth_total_rows(table, precio_unico, total_con_iva)
        
        # Aplicar estilos
        self.table_styles.style_gwealth_table(table)

    def _add_gwealth_total_rows(self, table, precio_unico: float, total_con_iva: float):
        """Agrega las filas de total para Gwealth."""
        # Fila TOTAL
        merged_total = self.table_styles.merge_row_cells(table, 2, 0, 1)
        merged_total.text = "TOTAL"
        for p in merged_total.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        table.cell(2, 2).text = format_currency(precio_unico)

        # Fila TOTAL CON IVA
        merged_total_iva = self.table_styles.merge_row_cells(table, 3, 0, 1)
        merged_total_iva.text = "TOTAL CON IVA"
        for p in merged_total_iva.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        table.cell(3, 2).text = format_currency(total_con_iva)
