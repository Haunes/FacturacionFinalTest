from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from io import BytesIO
import pandas as pd
from datetime import datetime
from utils.formatting_utils import format_currency
from utils.data_utils import get_document_count, get_representative_price
from .word_styles import WordStyleManager
from .word_table_builder import WordTableBuilder

class WordReportGenerator:
    """Generador de reportes en formato Word."""
    
    def __init__(self):
        self.style_manager = WordStyleManager()
        self.table_builder = WordTableBuilder()
    
    def generate_report(self, data: pd.DataFrame, empresa: str, anio: int, mes: str, funcionarios: dict) -> BytesIO:
        """
        Genera el documento Word completo.
        
        Args:
            data: Datos filtrados
            empresa: Nombre de la empresa
            anio: Año del reporte
            mes: Mes del reporte
            funcionarios: Información de funcionarios
            
        Returns:
            Buffer con el documento Word generado
        """
        doc = Document()
        
        # Configurar estilos del documento
        self.style_manager.setup_document_styles(doc)
        
        # Configurar márgenes
        self._setup_page_margins(doc)
        
        # Agregar contenido
        self._add_header(doc, mes, anio, empresa, funcionarios)
        self._add_main_table(doc, data, empresa)
        self._add_summary_tables(doc, data, empresa, anio, mes)
        self._add_footer(doc)
        
        # Guardar en buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _setup_page_margins(self, doc: Document):
        """Configura los márgenes de la página."""
        section = doc.sections[0]
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
    
    def _add_header(self, doc: Document, mes: str, anio: int, empresa: str, funcionarios: dict):
        """Agrega el encabezado del documento."""
        # Logo
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            p_logo.add_run().add_picture('assets/biu_logo.png', width=Inches(1.5))
        except FileNotFoundError:
            run = p_logo.add_run("[Logo BIU]")
            run.bold = True
        
        # Línea separadora
        p_rule = doc.add_paragraph()
        self.style_manager.set_paragraph_border_bottom(p_rule)
        
        # Títulos
        self._add_title(doc, f"FACTURACIÓN {mes.upper()} {anio}", size=24)
        self._add_title(doc, empresa.upper(), size=20)
        
        # Información del reporte
        doc.add_paragraph(" \nFecha de corte del reporte: ")
        doc.add_paragraph(f"Funcionario que reporta: \t {funcionarios['reporta']}")
        doc.add_paragraph(f"Funcionario revisor: \t\t {funcionarios['revisor']}")
    
    def _add_title(self, doc: Document, text: str, size: int):
        """Agrega un título con formato específico."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri Light'
        run.font.size = Pt(size)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)  # COLOR_PRIMARY
    
    def _add_main_table(self, doc: Document, data: pd.DataFrame, empresa: str):
        """Agrega la tabla principal de datos."""
        doc.add_paragraph()
        self.table_builder.add_main_table(doc, data, empresa)
    
    def _add_summary_tables(self, doc: Document, data: pd.DataFrame, empresa: str, anio: int, mes: str):
        """Agrega las tablas de resumen específicas por empresa."""
        doc.add_paragraph()
        self.table_builder.add_summary_tables(doc, data, empresa, anio, mes)
    
    def _add_footer(self, doc: Document):
        """Agrega el pie de página."""
        footer = doc.sections[0].footer
        p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p_footer.text = "Número: 601 - 7455289 | Dirección: Carrera 7 No. 74B-56, Oficina 301 | Correo: info@biu.com.co"
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in p_footer.runs:
            run.font.name = 'Calibri Light'
            run.font.size = Pt(9)
