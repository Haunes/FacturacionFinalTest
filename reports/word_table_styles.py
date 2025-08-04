from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class WordTableStyles:
    """Maneja los estilos de tablas para documentos Word."""
    
    # Colores
    COLOR_PRIMARY = RGBColor(0, 51, 102)
    COLOR_ACCENT = RGBColor(226, 0, 116)
    COLOR_LIGHT_GRAY = RGBColor(240, 240, 240)
    COLOR_BLACK = RGBColor(0, 0, 0)
    COLOR_WHITE = RGBColor(255, 255, 255)
    
    FONT_FAMILY = 'Calibri Light'
    
    def style_table(self, table, has_total_row=True):
        """Aplica estilos básicos a una tabla."""
        for i, row in enumerate(table.rows):
            is_header = (i == 0)
            is_total = has_total_row and (i == len(table.rows) - 1)
            
            for cell in row.cells:
                if is_header:
                    self.set_cell_shading(cell, self.COLOR_PRIMARY)
                elif is_total:
                    self.set_cell_shading(cell, self.COLOR_ACCENT)
                else:
                    self.set_cell_shading(cell, self.COLOR_LIGHT_GRAY)
                
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1
                    
                    for run in p.runs:
                        run.font.name = self.FONT_FAMILY
                        run.font.size = Pt(11)
                        
                        if is_header or is_total:
                            run.font.color.rgb = self.COLOR_WHITE
                            run.font.bold = True
                        else:
                            run.font.color.rgb = self.COLOR_BLACK
                            run.font.bold = False
    
    def style_gwealth_table(self, table):
        """Aplica estilos específicos para tablas de Gwealth."""
        self.style_table(table, has_total_row=False)
        
        # Filas 2 y 3 con color de acento
        for row_idx in [2, 3]:
            for cell in table.rows[row_idx].cells:
                self.set_cell_shading(cell, self.COLOR_ACCENT)
            self.set_row_text_color(table.rows[row_idx], self.COLOR_WHITE, bold=True)
        
        self.fix_table_layout_3cols(table)
        self.set_table_outer_borders(table)
        self.apply_3col_grid_borders(table, skip_mid_rows=(2, 3))
    
    def set_cell_shading(self, cell, color_rgb):
        """Establece el color de fondo de una celda."""
        hex_color = self._to_hex(color_rgb)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shd)
    
    def set_row_text_color(self, row, color_rgb=None, bold=True):
        """Establece el color de texto para toda una fila."""
        color_rgb = color_rgb or self.COLOR_WHITE
        
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = color_rgb
                    if bold is not None:
                        run.bold = bold
    
    def merge_row_cells(self, table, row_idx: int, start_col: int, end_col: int):
        """Fusiona celdas en una fila."""
        top_left = table.cell(row_idx, start_col)
        bottom_right = table.cell(row_idx, end_col)
        return top_left.merge(bottom_right)
    
    def fix_table_layout_3cols(self, table, widths=(Inches(1.2), Inches(4.2), Inches(1.4))):
        """Ajusta el layout para tablas de 3 columnas."""
        table.autofit = False
        
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths):
                    cell.width = widths[idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1
            
            row.height = Inches(0.28)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    def set_table_borders(self, table, color="000000", size=8, val="single"):
        """Establece bordes para toda la tabla."""
        tbl = table._tbl
        tblPr = tbl.tblPr or OxmlElement('w:tblPr')
        if tbl.tblPr is None:
            tbl.append(tblPr)
        
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        
        tblBorders = OxmlElement('w:tblBorders')
        for edge in ('w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV'):
            el = OxmlElement(edge)
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), str(int(size)))
            el.set(qn('w:color'), color)
            el.set(qn('w:space'), '0')
            tblBorders.append(el)
        
        tblPr.append(tblBorders)
    
    def set_table_outer_borders(self, table, color="000000", size=8, val="single"):
        """Establece solo los bordes exteriores de la tabla."""
        tbl = table._tbl
        tblPr = tbl.tblPr or OxmlElement('w:tblPr')
        if tbl.tblPr is None:
            tbl.append(tblPr)
        
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        
        tblBorders = OxmlElement('w:tblBorders')
        for edge in ('w:top', 'w:left', 'w:bottom', 'w:right'):
            el = OxmlElement(edge)
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), str(int(size)))
            el.set(qn('w:color'), color)
            el.set(qn('w:space'), '0')
            tblBorders.append(el)
        
        tblPr.append(tblBorders)
    
    def apply_3col_grid_borders(self, table, skip_mid_rows=()):
        """Aplica bordes específicos para tablas de 3 columnas."""
        # Implementación simplificada - en un caso real sería más compleja
        pass
    
    def _to_hex(self, color) -> str:
        """Convierte un color RGB a hexadecimal."""
        if isinstance(color, RGBColor):
            return str(color)
        if isinstance(color, (tuple, list)) and len(color) == 3:
            r, g, b = map(int, color)
            return f"{r:02x}{g:02x}{b:02x}"
        if isinstance(color, str):
            s = color.lstrip("#")
            return s.upper() if len(s) == 6 else "FFFFFF"
        return "FFFFFF"
